"""
exam_docx_template.py
--------------------
Create a compact, two-column A4-landscape exam paper DOCX (Times New Roman),
matching the style you showed: very low vertical spacing, small left padding,
headers slightly larger (NOT bold), only the very top title bold.

INPUT FORMAT (plain text):
- Prefix each line with one of these markers:

T:  Main title (bold, centered, full-width header section)
M:  Meta line (full-width header section). Use tabs for alignment:
    e.g. "M: Sub:Science\tFM: 50"
S:  Section header (slightly larger, not bold)
Q:  Question line (small left indent)
O:  Option/inner list line (more indent)
R:  Rule line (e.g. stars)
P:  Plain paragraph line (body text)

Example input.txt:
T: Third Terminal Examination 2082
M: Sub:Science\tFM: 50
M: Class: 6\tTime: 1.30 hrs\tPM: 20
R: ******************************************************
S: 1. Multiple Choice Questions. (10*1=10)
Q: a. Where do we live?
O: i. crust\t ii. mantle\t iii. outer core\t iv. inner core
Q: b. Which of the following is a magnetic substances?
O: i. Iron\t ii. wood\t iii. plastic\t iv. rubber
S: 2. Very short answer questions. (10*1=10)
Q: a. Define top soil.

USAGE:
  pip install python-docx
  python exam_docx_template.py input.txt output.docx
"""

import sys
from dataclasses import dataclass
from typing import List, Tuple, Optional

from docx import Document
from docx.shared import Pt, Cm, Length
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ---------- CONFIG (tweak these to match your “house style”) ----------
FONT_NAME = "Times New Roman"

# Page / columns
PAGE_WIDTH_CM = 29.7   # A4 landscape
PAGE_HEIGHT_CM = 21.0
MARGIN_CM = 1.0
COL_GAP_TWIPS = 260    # column gap (twips) for a compact look
COLUMN_SEPARATOR = True

# Compact spacing
BASE_FONT_PT = 13
BASE_LINE_PT = 15      # Increased from 12 for better line spacing when text wraps
HEADER_TITLE_PT = 16
META_PT = 14
SECTION_PT = 14

# Vertical spacing (space before/after paragraphs in points)
SPACE_BEFORE_TITLE = 0
SPACE_AFTER_TITLE = 6
SPACE_BEFORE_META = 0
SPACE_AFTER_META = 3
SPACE_BEFORE_SECTION = 6
SPACE_AFTER_SECTION = 3
SPACE_BEFORE_QUESTION = 5  # Moderate increase from 3 (half of 8)
SPACE_AFTER_QUESTION = 2   # Moderate increase from 0 (half of 4)
SPACE_BEFORE_OPTION = 2    # Moderate spacing before options
SPACE_AFTER_OPTION = 1     # Moderate spacing after options
SPACE_BEFORE_PARAGRAPH = 0
SPACE_AFTER_PARAGRAPH = 0

# Indents (compact)
Q_LEFT_CM = 0.4
O_LEFT_CM = 0.8

# If True, section headers are kept with the next paragraph
KEEP_SECTION_WITH_NEXT = True


# ---------- Simple model ----------
@dataclass
class Block:
    kind: str   # "title", "meta", "section", "question", "option", "rule", "para"
    text: str


# ---------- Helpers: Word XML for columns ----------
def _set_columns(section, num: int, space_twips: int, sep: bool) -> None:
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))
    if sep:
        cols.set(qn("w:sep"), "1")
    else:
        if qn("w:sep") in cols.attrib:
            del cols.attrib[qn("w:sep")]


def _set_style_font(style, font_name: str) -> None:
    """Force font for all char sets."""
    style.font.name = font_name
    rPr = style._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def _setup_page(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)

    section.top_margin = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)
    section.left_margin = Cm(MARGIN_CM)
    section.right_margin = Cm(MARGIN_CM)


def _usable_width(section) -> Length:
    return section.page_width - section.left_margin - section.right_margin


def _column_width_twips(section) -> int:
    """
    Calculate the width of one column in twips.
    """
    # usable page width = page - margins (in twips)
    # Convert each Length to twips before subtraction
    page_width_twips = int(section.page_width.twips)
    left_margin_twips = int(section.left_margin.twips)
    right_margin_twips = int(section.right_margin.twips)
    usable_twips = page_width_twips - left_margin_twips - right_margin_twips

    # gap between columns is already in twips
    gap_twips = COL_GAP_TWIPS

    # width of one column in twips
    return (usable_twips - gap_twips) // 2


def _ensure_styles(doc: Document) -> None:
    # Base Normal
    normal = doc.styles["Normal"]
    _set_style_font(normal, FONT_NAME)
    normal.font.size = Pt(BASE_FONT_PT)
    normal.font.bold = False
    pf = normal.paragraph_format
    pf.space_before = Pt(SPACE_BEFORE_PARAGRAPH)
    pf.space_after = Pt(SPACE_AFTER_PARAGRAPH)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(BASE_LINE_PT)

    def add_para_style(name: str) -> None:
        if name in doc.styles:
            return
        doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    # Title
    add_para_style("ExamTitle")
    st = doc.styles["ExamTitle"]
    _set_style_font(st, FONT_NAME)
    st.font.size = Pt(HEADER_TITLE_PT)
    st.font.bold = True
    st.paragraph_format.space_before = Pt(SPACE_BEFORE_TITLE)
    st.paragraph_format.space_after = Pt(SPACE_AFTER_TITLE)
    st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    st.paragraph_format.line_spacing = Pt(HEADER_TITLE_PT + 2)

    # Meta
    add_para_style("ExamMeta")
    st = doc.styles["ExamMeta"]
    _set_style_font(st, FONT_NAME)
    st.font.size = Pt(META_PT)
    st.font.bold = False
    st.paragraph_format.space_before = Pt(SPACE_BEFORE_META)
    st.paragraph_format.space_after = Pt(SPACE_AFTER_META)
    st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    st.paragraph_format.line_spacing = Pt(META_PT + 1)

    # Section header (slightly bigger, NOT bold)
    add_para_style("ExamSection")
    st = doc.styles["ExamSection"]
    _set_style_font(st, FONT_NAME)
    st.font.size = Pt(SECTION_PT)
    st.font.bold = False
    st.paragraph_format.space_before = Pt(SPACE_BEFORE_SECTION)
    st.paragraph_format.space_after = Pt(SPACE_AFTER_SECTION)
    st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    st.paragraph_format.line_spacing = Pt(SECTION_PT + 1)

    # Question
    add_para_style("ExamQuestion")
    st = doc.styles["ExamQuestion"]
    _set_style_font(st, FONT_NAME)
    st.font.size = Pt(BASE_FONT_PT)
    st.font.bold = False
    st.paragraph_format.left_indent = Cm(Q_LEFT_CM)
    st.paragraph_format.space_before = Pt(SPACE_BEFORE_QUESTION)
    st.paragraph_format.space_after = Pt(SPACE_AFTER_QUESTION)
    st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    st.paragraph_format.line_spacing = Pt(BASE_LINE_PT)

    # Option / inner list
    add_para_style("ExamOption")
    st = doc.styles["ExamOption"]
    _set_style_font(st, FONT_NAME)
    st.font.size = Pt(BASE_FONT_PT)
    st.font.bold = False
    st.paragraph_format.left_indent = Cm(O_LEFT_CM)
    st.paragraph_format.space_before = Pt(SPACE_BEFORE_OPTION)
    st.paragraph_format.space_after = Pt(SPACE_AFTER_OPTION)
    st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    st.paragraph_format.line_spacing = Pt(BASE_LINE_PT)


def parse_blocks(text: str) -> List[Block]:
    blocks: List[Block] = []
    for raw in text.splitlines():
        line = raw.rstrip("\n")
        if not line.strip():
            # for compact look, ignore blank lines
            continue

        # allow writing "\t" in the file (instead of a real tab)
        line = line.replace("\\t", "\t")

        # detect "X: " prefix
        if len(line) >= 2 and line[1] == ":":
            tag = line[0].upper()
            content = line[2:].lstrip()
        else:
            # default: treat as paragraph
            tag = "P"
            content = line.strip()

        kind_map = {
            "T": "title",
            "M": "meta",
            "S": "section",
            "Q": "question",
            "O": "option",
            "R": "rule",
            "P": "para",
        }
        kind = kind_map.get(tag, "para")
        blocks.append(Block(kind=kind, text=content))
    return blocks


def _add_meta_paragraph(doc: Document, section, text: str) -> None:
    """
    Meta lines placed INSIDE left column, with tab stops relative to column width.
    Example: 'Sub:Science\tFM: 50'
    """
    p = doc.add_paragraph(text, style="ExamMeta")

    col_w_twips = _column_width_twips(section)
    p.paragraph_format.tab_stops.clear_all()
    p.paragraph_format.tab_stops.add_tab_stop(col_w_twips // 2, alignment=WD_TAB_ALIGNMENT.CENTER)
    p.paragraph_format.tab_stops.add_tab_stop(col_w_twips, alignment=WD_TAB_ALIGNMENT.RIGHT)

    # keeps the header block together in the left column (nice compact look)
    p.paragraph_format.keep_with_next = True


def build_docx_from_blocks(blocks: List[Block], out_path: str) -> None:
    doc = Document()
    _ensure_styles(doc)

    # Start in 2 columns immediately (so header stays in LEFT column)
    sec = doc.sections[0]
    _setup_page(sec)
    _set_columns(sec, num=2, space_twips=COL_GAP_TWIPS, sep=COLUMN_SEPARATOR)

    for idx, b in enumerate(blocks):
        if b.kind == "title":
            p = doc.add_paragraph(b.text, style="ExamTitle")
            # centered within left column; change to LEFT if you prefer:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True

        elif b.kind == "meta":
            _add_meta_paragraph(doc, sec, b.text)

        elif b.kind == "rule":
            p = doc.add_paragraph(b.text, style="Normal")
            p.paragraph_format.keep_with_next = True

        elif b.kind == "section":
            p = doc.add_paragraph(b.text, style="ExamSection")
            if KEEP_SECTION_WITH_NEXT:
                p.paragraph_format.keep_with_next = True

        elif b.kind == "question":
            doc.add_paragraph(b.text, style="ExamQuestion")

        elif b.kind == "option":
            doc.add_paragraph(b.text, style="ExamOption")

        else:
            doc.add_paragraph(b.text, style="Normal")

    doc.save(out_path)


def main():
    if len(sys.argv) < 3:
        print("Usage: python exam_docx_template.py input.txt output.docx")
        sys.exit(1)

    in_path = sys.argv[1]
    out_path = sys.argv[2]

    with open(in_path, "r", encoding="utf-8") as f:
        text = f.read()

    blocks = parse_blocks(text)
    build_docx_from_blocks(blocks, out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
