"""
Exam DOCX generator — matches reference files EXACTLY.

Reference approach:
- Normal style only (no custom styles)
- Leading spaces for indentation
- Line spacing: SINGLE (1.0) for landscape, ONE_POINT_FIVE (1.5) for portrait
- Bold on runs for title, meta, question numbers
- Tabs for alignment on meta lines
"""
from io import BytesIO
from docx import Document
from docx.shared import Pt, Twips
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import Optional, List, Union

try:
    from exam_schema import ExamDocument
except ImportError:
    ExamDocument = None

FONT = "Times New Roman"


# ============================================================
# Configs — from reference DOCX inspection
# ============================================================
LANDSCAPE = {
    "orientation": WD_ORIENT.LANDSCAPE,
    "page_w_twips": 16839,
    "page_h_twips": 11907,
    "margin_t": 432, "margin_b": 432,
    "margin_l": 720, "margin_r": 720,
    "cols": 2, "col_space": 720, "col_sep": True,
    "font_size_pt": 14,
    "title_font_size_pt": 16,
    "line_spacing_rule": WD_LINE_SPACING.SINGLE,
    "line_spacing_value": 1.0,
    "meta_bold": True,
}

PORTRAIT = {
    "orientation": WD_ORIENT.PORTRAIT,
    "page_w_twips": 11909,
    "page_h_twips": 16834,
    "margin_t": 720, "margin_b": 720,
    "margin_l": 720, "margin_r": 720,
    "cols": 1, "col_space": 720, "col_sep": False,
    "font_size_pt": 18,
    "title_font_size_pt": 18,
    "line_spacing_rule": WD_LINE_SPACING.ONE_POINT_FIVE,
    "line_spacing_value": 1.5,
    "meta_bold": True,
}


# ============================================================
# Helpers
# ============================================================

def _set_run_font(run, font_name: str, size_pt: int, bold: bool = False):
    """Set font on a run, forcing all character sets."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold if bold else None
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def _set_columns(section, num, space, sep):
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space))
    if sep:
        cols.set(qn("w:sep"), "1")


def _set_normal_font(doc, size_pt):
    """Set Times New Roman at the Normal style level."""
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(size_pt)
    rPr = normal._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:eastAsia"), FONT)
    rFonts.set(qn("w:cs"), FONT)


def _add_paragraph(doc, text: str, bold: bool = False, center: bool = False,
                   line_spacing=None, keep_next=False):
    """
    Add a paragraph using Normal style.
    Set font/bold at run level, matching the reference approach.
    """
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, FONT, 14, bold)  # default 14, overridden below

    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if keep_next:
        p.paragraph_format.keep_with_next = True

    return p


def _add_paragraph_sized(doc, text: str, size_pt: int, bold: bool = False,
                         center: bool = False, keep_next=False,
                         line_spacing_rule=None, line_spacing_value=None):
    """
    Add paragraph with explicit font size.
    All formatting at run level, matching reference exactly.
    """
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, FONT, size_pt, bold)

    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if keep_next:
        p.paragraph_format.keep_with_next = True
    if line_spacing_rule is not None:
        p.paragraph_format.line_spacing_rule = line_spacing_rule
        p.paragraph_format.line_spacing = line_spacing_value

    return p


# ============================================================
# ExamDocument → DOCX
# ============================================================

def build_exam_docx(
    output_path: Union[str, BytesIO],
    exam_data: Optional['ExamDocument'] = None,
    content: Optional[List[tuple]] = None,
    orientation: str = "landscape",
):
    """Build exam DOCX matching reference formatting exactly."""
    config = LANDSCAPE if orientation == "landscape" else PORTRAIT
    fsz = config["font_size_pt"]
    tfsz = config["title_font_size_pt"]
    ls_rule = config["line_spacing_rule"]
    ls_val = config["line_spacing_value"]
    bold_meta = config["meta_bold"]

    doc = Document()

    # Page setup
    sec = doc.sections[0]
    sec.orientation = config["orientation"]
    sec.page_width = Twips(config["page_w_twips"])
    sec.page_height = Twips(config["page_h_twips"])
    sec.top_margin = Twips(config["margin_t"])
    sec.bottom_margin = Twips(config["margin_b"])
    sec.left_margin = Twips(config["margin_l"])
    sec.right_margin = Twips(config["margin_r"])
    _set_columns(sec, config["cols"], config["col_space"], config["col_sep"])

    # Normal style font
    _set_normal_font(doc, fsz)

    # --- Build paragraphs ---
    if exam_data:
        _write_exam(doc, exam_data, orientation, config)
    elif content:
        _write_legacy(doc, content, config)

    doc.save(output_path)


def _write_exam(doc, exam: 'ExamDocument', orientation: str, config: dict):
    """Write ExamDocument content."""
    fsz = config["font_size_pt"]
    tfsz = config["title_font_size_pt"]
    ls_rule = config["line_spacing_rule"]
    ls_val = config["line_spacing_value"]
    bold_meta = config["meta_bold"]

    # Title
    if exam.title:
        _add_paragraph_sized(doc, exam.title, tfsz, bold=True, center=True,
                             keep_next=True, line_spacing_rule=ls_rule, line_spacing_value=ls_val)

    # Metadata lines
    meta_lines = exam.metadata.to_meta_lines()
    for i, line in enumerate(meta_lines):
        p = _add_paragraph_sized(doc, line, fsz, bold=bold_meta, center=True,
                                 keep_next=True, line_spacing_rule=ls_rule, line_spacing_value=ls_val)

    # Instructions
    if exam.instructions:
        _add_paragraph_sized(doc, exam.instructions, fsz, center=True,
                             keep_next=True, line_spacing_rule=ls_rule, line_spacing_value=ls_val)

    # Separator
    _add_paragraph_sized(doc, '*' * 50, fsz, keep_next=True,
                         line_spacing_rule=ls_rule, line_spacing_value=ls_val)

    # Sections & questions
    for section in exam.sections:
        header = section.format_header()
        if header:
            _add_paragraph_sized(doc, header, fsz, bold=False, keep_next=True,
                                 line_spacing_rule=ls_rule, line_spacing_value=ls_val)

        for question in section.questions:
            q_text = question.format_question()
            _add_paragraph_sized(doc, q_text, fsz, line_spacing_rule=ls_rule,
                                 line_spacing_value=ls_val)

            if question.sub_questions:
                for sub_q in question.sub_questions:
                    # Sub-question with leading spaces (8 spaces like reference)
                    sq_text = "        " + sub_q.format_question()
                    _add_paragraph_sized(doc, sq_text, fsz, line_spacing_rule=ls_rule,
                                         line_spacing_value=ls_val)
                    if sub_q.options:
                        opt_text = sub_q.format_options()
                        if opt_text:
                            # Options with more leading spaces (11 spaces like reference)
                            _add_paragraph_sized(doc, "           " + opt_text, fsz,
                                                 line_spacing_rule=ls_rule,
                                                 line_spacing_value=ls_val)

            if question.options:
                opt_text = question.format_options()
                if opt_text:
                    _add_paragraph_sized(doc, "           " + opt_text, fsz,
                                         line_spacing_rule=ls_rule,
                                         line_spacing_value=ls_val)


def _write_legacy(doc, content: List[tuple], config: dict):
    """Write legacy tagged text content."""
    fsz = config["font_size_pt"]
    tfsz = config["title_font_size_pt"]
    ls_rule = config["line_spacing_rule"]
    ls_val = config["line_spacing_value"]
    bold_meta = config["meta_bold"]

    for text, style_name, alignment, tabs in content:
        if style_name == "ExamTitle":
            _add_paragraph_sized(doc, text, tfsz, bold=True, center=True,
                                 keep_next=True, line_spacing_rule=ls_rule, line_spacing_value=ls_val)
        elif style_name == "ExamMeta":
            _add_paragraph_sized(doc, text, fsz, bold=bold_meta, center=True,
                                 keep_next=True, line_spacing_rule=ls_rule, line_spacing_value=ls_val)
        elif style_name == "ExamSectionHeader":
            _add_paragraph_sized(doc, text, fsz, keep_next=True,
                                 line_spacing_rule=ls_rule, line_spacing_value=ls_val)
        elif style_name == "ExamQuestion":
            _add_paragraph_sized(doc, text, fsz, line_spacing_rule=ls_rule,
                                 line_spacing_value=ls_val)
        elif style_name == "ExamOption":
            _add_paragraph_sized(doc, "           " + text, fsz, line_spacing_rule=ls_rule,
                                 line_spacing_value=ls_val)
        else:
            _add_paragraph_sized(doc, text, fsz, line_spacing_rule=ls_rule,
                                 line_spacing_value=ls_val)


# ============================================================
# Legacy: parse tagged text format
# ============================================================

def parse_blocks(text: str):
    from dataclasses import dataclass
    @dataclass
    class Block:
        kind: str
        text: str

    blocks = []
    for raw in text.splitlines():
        line = raw.rstrip("\n")
        if not line.strip():
            continue
        line = line.replace("\\t", "\t")
        if len(line) >= 2 and line[1] == ":":
            tag = line[0].upper()
            content = line[2:].lstrip()
        else:
            tag = "P"
            content = line.strip()

        kind_map = {
            "T": "title", "M": "meta", "S": "section",
            "Q": "question", "O": "option", "R": "rule", "P": "para",
        }
        blocks.append(Block(kind=kind_map.get(tag, "para"), text=content))
    return blocks


def build_docx_from_text(text: str, output_path: Union[str, BytesIO], orientation: str = "landscape"):
    """Build DOCX from tagged text format."""
    blocks = parse_blocks(text)
    config = LANDSCAPE if orientation == "landscape" else PORTRAIT

    doc = Document()
    sec = doc.sections[0]
    sec.orientation = config["orientation"]
    sec.page_width = Twips(config["page_w_twips"])
    sec.page_height = Twips(config["page_h_twips"])
    sec.top_margin = Twips(config["margin_t"])
    sec.bottom_margin = Twips(config["margin_b"])
    sec.left_margin = Twips(config["margin_l"])
    sec.right_margin = Twips(config["margin_r"])
    _set_columns(sec, config["cols"], config["col_space"], config["col_sep"])
    _set_normal_font(doc, config["font_size_pt"])

    fsz = config["font_size_pt"]
    tfsz = config["title_font_size_pt"]
    ls_rule = config["line_spacing_rule"]
    ls_val = config["line_spacing_value"]
    bold_meta = config["meta_bold"]

    for b in blocks:
        if b.kind == "title":
            _add_paragraph_sized(doc, b.text, tfsz, bold=True, center=True,
                                 keep_next=True, line_spacing_rule=ls_rule, line_spacing_value=ls_val)
        elif b.kind == "meta":
            _add_paragraph_sized(doc, b.text, fsz, bold=bold_meta, center=True,
                                 keep_next=True, line_spacing_rule=ls_rule, line_spacing_value=ls_val)
        elif b.kind == "section":
            _add_paragraph_sized(doc, b.text, fsz, keep_next=True,
                                 line_spacing_rule=ls_rule, line_spacing_value=ls_val)
        elif b.kind == "question":
            _add_paragraph_sized(doc, b.text, fsz, line_spacing_rule=ls_rule,
                                 line_spacing_value=ls_val)
        elif b.kind == "option":
            _add_paragraph_sized(doc, "           " + b.text, fsz, line_spacing_rule=ls_rule,
                                 line_spacing_value=ls_val)
        elif b.kind == "rule":
            _add_paragraph_sized(doc, b.text, fsz, keep_next=True,
                                 line_spacing_rule=ls_rule, line_spacing_value=ls_val)
        else:
            _add_paragraph_sized(doc, b.text, fsz, line_spacing_rule=ls_rule,
                                 line_spacing_value=ls_val)

    doc.save(output_path)


if __name__ == "__main__":
    build_docx_from_text(
        open("oldinput.txt").read(),
        "demo_landscape.docx",
        orientation="landscape",
    )
    build_docx_from_text(
        open("oldinput.txt").read(),
        "demo_portrait.docx",
        orientation="portrait",
    )
    print("Saved demo_landscape.docx and demo_portrait.docx")
